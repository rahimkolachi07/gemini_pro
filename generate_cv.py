from docx import Document

def generate_cv_template(name, cv_summary, cv_skills):
    # CV template
    cv_template = f"""
    ======= CV Template =======
    **Name:** {name}

    **Summary:**
    {cv_summary}

    **Skills:**
    {cv_skills}
    ==========================
    """
    doc = Document()

    # Add content to the Word document
    doc.add_heading('Curriculum Vitae', level = 1).bold = True

    # Add name
    doc.add_heading(name, level=2).bold = True

    # Add summary
    doc.add_heading('Summary', level=2).bold = True
    doc.add_paragraph(cv_summary)

    # Add Skills
    doc.add_heading('Skills', level=2).bold = True
    doc.add_paragraph(cv_skills)

    # Save the Word document
    doc.save(f'D:\\HectoByte\\gemini_pro\\{name.replace(" ", "_")}_cv.docx')
